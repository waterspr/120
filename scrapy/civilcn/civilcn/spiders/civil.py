# -*- coding: utf-8 -*-
import scrapy
import requests
from docx import Document
from docx.shared import Inches


class CivilSpider(scrapy.Spider):
    name = 'civil'
    allowed_domains = ['civilcn.com']
    start_urls = ['http://www.civilcn.com/zhishi/sgjs/']

    def parse(self, response):
        # 获取内层的文章的链接
        urllists = response.xpath(
            '//div[@class="m_g_b_d"]//li/span/a/@href').extract()  # 返回是list列表
        # print(urllists)
        for url in urllists:
            print(url)
            yield scrapy.Request(url, callback=self.parse2)  # 内层网页解析
            # 测试二层网页解析
            # yield scrapy.Request('http://www.civilcn.com/zhishi/sgjs/1572588502379220.html', callback=self.parse2)  #内层网页解析

        # 构造下一页url
        next_text = response.xpath(
            '//div[@class="m_g_b_c"]/a[last()-1]/text()').extract_first()
        if next_text == '下一页':
            next_page = response.xpath(
                '//div[@class="m_g_b_c"]/a[last()-1]/@href').extract_first()
            print('-'*35)
            print(next_page)
            yield scrapy.Request(next_page, callback=self.parse)

    def parse2(self, response):
        title = response.xpath(
            '//div[@class="m_g_title"]/text()').extract_first()
        print(title)

        document = Document()  # 新建word文档
        document.add_heading(title)

        # 按段落取数据
        ps = response.xpath('//div[@class="m_g_content"]/p')
        for p in ps:

            # 判断是否有图片
            img = p.xpath('.//img/@src').extract_first()
            if img:
                img_name = img[-9:]  # 指定图片的名
                print(img_name)
                with open('f:/1/'+img_name, 'wb') as fp:
                    fp.write(requests.get(img).content)
                try:
                    document.add_picture(
                        'f:/1/'+img_name, width=Inches(6))  # 写入图片到word
                except:
                    pass
            # 判断二级标题  # 一个段落里面既有加粗，也有文本
            strong_name = p.xpath('./strong/text()').extract_first()
            if strong_name:
                document.add_heading(strong_name, level=2)
                # 取文本
                content1 = p.xpath('./text()').extract_first()  # 返回是个字符串
                # print(content1)
                try:
                    content1 = content1.replace('\n', '').replace(
                        '\r', '')  # 去掉回车符\r和换行符\n两种标记
                except:
                    pass
                document.add_paragraph(content1)
            # #提取文本内容

            # extract_first()返回是第一个字符串，取不全
            contents = p.xpath('./text()').extract()
            for content in contents:
                # print(content)
                try:
                    content = content.replace('\n', '').replace(
                        '\r', '')  # 去掉回车符\r和换行符\n两种标记
                except:
                    pass
                document.add_paragraph(content)

        document.save('f:/2/'+title + '.docx')
